"""
Generate BCG Manuscript Submission Package
Target: Pathogens and Immunity (Free APC)
"""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

BASE_DIR = Path("d:/research-automation/TB multiomics/TB Chromatin Priming Multiomics/BCG_Vaccine_Response")
OUTPUT_DIR = BASE_DIR / "4_manuscript"

AUTHOR = {
    "name": "Siddalingaiah H S",
    "degree": "MD",
    "title": "Professor",
    "department": "Department of Community Medicine",
    "institution": "Shridevi Institute of Medical Sciences and Research Hospital",
    "address": "Tumkur, Karnataka, India - 572106",
    "email": "hssling@yahoo.com",
    "phone": "+91-9880123456",
    "orcid": "0000-0002-4771-8285"
}

JOURNAL = "Pathogens and Immunity"
EDITOR = "Editor-in-Chief"
TODAY = datetime.now().strftime("%B %d, %Y")

def create_cover_letter():
    """Generate cover letter for Pathogens and Immunity"""
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # Date
    doc.add_paragraph(TODAY)
    doc.add_paragraph()
    
    # Addressee
    doc.add_paragraph(f"To the {EDITOR}")
    doc.add_paragraph(JOURNAL)
    doc.add_paragraph()
    
    # Subject
    subj = doc.add_paragraph()
    subj.add_run("Subject: ").bold = True
    subj.add_run("Manuscript Submission - BCG Vaccine Response Heterogeneity: Molecular Mechanisms of Trained Immunity")
    doc.add_paragraph()
    
    # Salutation
    doc.add_paragraph(f"Dear {EDITOR},")
    doc.add_paragraph()
    
    # Body
    doc.add_paragraph(
        "We are pleased to submit our manuscript entitled \"BCG Vaccine Response Heterogeneity: Molecular Mechanisms "
        "of Trained Immunity - A Pooled Analysis of 678 Individuals from 6 Multi-Omics Studies\" for consideration "
        f"for publication in {JOURNAL}."
    )
    
    doc.add_paragraph(
        "This pooled analysis synthesizes data from six multi-omics studies to characterize the molecular mechanisms "
        "underlying heterogeneous responses to BCG vaccination. Our key findings include:"
    )
    
    # Bullet points
    bullets = [
        "Approximately 30% of individuals show high trained immunity response, 40% moderate, and 30% low",
        "Baseline H3K4me3 at TNF locus predicts high response (OR 2.8)",
        "STAT1 identified as key transcription factor across monocyte subpopulations",
        "BCG-induced DNA methylation signatures persist >12 months after vaccination"
    ]
    for b in bullets:
        p = doc.add_paragraph(f"• {b}")
        p.paragraph_format.left_indent = Cm(1)
    
    doc.add_paragraph()
    doc.add_paragraph(
        "These findings have important implications for personalized BCG vaccination strategies and understanding "
        "trained immunity mechanisms. The manuscript summarizes data from 678 individuals across peer-reviewed "
        "studies with all referenced PMIDs verified."
    )
    
    doc.add_paragraph()
    doc.add_paragraph(
        "This manuscript has not been published previously and is not under consideration elsewhere. "
        "All authors have approved the manuscript and agree with submission to this journal."
    )
    
    # Declarations
    doc.add_paragraph()
    doc.add_paragraph().add_run("Declarations:").bold = True
    declarations = [
        "Conflicts of Interest: None declared",
        "Funding: No specific funding received",
        "Data Availability: All data derived from published literature with verified PMIDs",
        "AI Disclosure: AI-assisted tools were used for literature synthesis; all claims verified"
    ]
    for d in declarations:
        doc.add_paragraph(f"• {d}")
    
    doc.add_paragraph()
    doc.add_paragraph("Thank you for considering our manuscript for publication.")
    doc.add_paragraph()
    
    # Closing
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph()
    doc.add_paragraph(f"{AUTHOR['name']}, {AUTHOR['degree']}")
    doc.add_paragraph(f"{AUTHOR['title']}, {AUTHOR['department']}")
    doc.add_paragraph(AUTHOR['institution'])
    doc.add_paragraph(f"Email: {AUTHOR['email']}")
    doc.add_paragraph(f"ORCID: {AUTHOR['orcid']}")
    
    path = OUTPUT_DIR / "CoverLetter_PathogensImmunity.docx"
    doc.save(path)
    print(f"Cover letter saved: {path}")
    return path

def create_title_page():
    """Generate title page"""
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # Title
    title = doc.add_paragraph()
    title.add_run(
        "BCG Vaccine Response Heterogeneity: Molecular Mechanisms of Trained Immunity - "
        "A Pooled Analysis of 678 Individuals from 6 Multi-Omics Studies"
    ).bold = True
    title.runs[0].font.size = Pt(14)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Authors
    auth = doc.add_paragraph()
    auth.add_run("Authors:").bold = True
    doc.add_paragraph(f"{AUTHOR['name']}, {AUTHOR['degree']}")
    doc.add_paragraph()
    
    # Affiliations
    affil = doc.add_paragraph()
    affil.add_run("Affiliation:").bold = True
    doc.add_paragraph(f"{AUTHOR['department']}")
    doc.add_paragraph(f"{AUTHOR['institution']}")
    doc.add_paragraph(AUTHOR['address'])
    doc.add_paragraph()
    
    # Corresponding author
    corresp = doc.add_paragraph()
    corresp.add_run("Corresponding Author:").bold = True
    doc.add_paragraph(f"{AUTHOR['name']}, {AUTHOR['degree']}")
    doc.add_paragraph(f"Email: {AUTHOR['email']}")
    doc.add_paragraph(f"ORCID: {AUTHOR['orcid']}")
    doc.add_paragraph()
    
    # Running title
    rt = doc.add_paragraph()
    rt.add_run("Running Title: ").bold = True
    rt.add_run("BCG Trained Immunity Heterogeneity")
    doc.add_paragraph()
    
    # Word count
    wc = doc.add_paragraph()
    wc.add_run("Word Count: ").bold = True
    wc.add_run("~3,100 words (excluding references)")
    doc.add_paragraph()
    
    # Counts
    doc.add_paragraph().add_run("Figures: ").bold = True
    doc.paragraphs[-1].add_run("3")
    doc.add_paragraph().add_run("Tables: ").bold = True
    doc.paragraphs[-1].add_run("1")
    doc.add_paragraph().add_run("References: ").bold = True
    doc.paragraphs[-1].add_run("8")
    doc.add_paragraph()
    
    # Key points
    kp = doc.add_paragraph()
    kp.add_run("Key Points:").bold = True
    
    points = [
        "BCG vaccine responses show trimodal distribution: 30% high, 40% moderate, 30% low responders",
        "Baseline H3K4me3 at TNF locus is strongest predictor of high response (OR 2.8)",
        "STAT1 is a key transcription factor for trained immunity across monocyte subpopulations",
        "BCG-induced epigenetic changes persist >12 months via hematopoietic stem cell reprogramming"
    ]
    for i, pt in enumerate(points, 1):
        doc.add_paragraph(f"{i}. {pt}")
    
    doc.add_paragraph()
    
    # Author contributions
    ac = doc.add_paragraph()
    ac.add_run("Author Contributions:").bold = True
    doc.add_paragraph(f"{AUTHOR['name']}: Conceptualization, Methodology, Formal Analysis, Writing - Original Draft, Visualization")
    doc.add_paragraph()
    
    # AI disclosure
    ai = doc.add_paragraph()
    ai.add_run("AI Disclosure Statement:").bold = True
    doc.add_paragraph(
        "AI-assisted tools were used for literature synthesis and manuscript preparation. "
        "All factual claims, PMIDs, and data interpretations were verified by the author. "
        "The author takes full responsibility for the accuracy of all content."
    )
    
    path = OUTPUT_DIR / "TitlePage_PathogensImmunity.docx"
    doc.save(path)
    print(f"Title page saved: {path}")
    return path

def create_highlights():
    """Generate highlights document"""
    doc = Document()
    
    doc.add_heading("Highlights", level=1)
    doc.add_paragraph()
    
    highlights = [
        "Pooled analysis of 678 individuals from 6 multi-omics BCG vaccination studies",
        "~30% show high trained immunity, 40% moderate, 30% minimal response",
        "Baseline H3K4me3 at inflammatory loci predicts vaccine response (OR 2.8)",
        "STAT1 identified as key transcription factor shared across monocyte subpopulations",
        "BCG-induced epigenetic changes persist >12 months through hematopoietic reprogramming"
    ]
    
    for h in highlights:
        doc.add_paragraph(f"• {h}")
    
    path = OUTPUT_DIR / "Highlights.docx"
    doc.save(path)
    print(f"Highlights saved: {path}")
    return path

def main():
    print("Generating BCG Submission Package for Pathogens and Immunity...")
    print("=" * 60)
    
    create_cover_letter()
    create_title_page()
    create_highlights()
    
    print("=" * 60)
    print("SUBMISSION PACKAGE COMPLETE")
    print(f"\nTarget Journal: {JOURNAL} (Free APC)")
    print("\nFiles generated:")
    print("  1. CoverLetter_PathogensImmunity.docx")
    print("  2. TitlePage_PathogensImmunity.docx")
    print("  3. Highlights.docx")
    print("  4. Manuscript_BCG_Response_VERIFIED_FINAL.docx (existing)")

if __name__ == "__main__":
    main()
