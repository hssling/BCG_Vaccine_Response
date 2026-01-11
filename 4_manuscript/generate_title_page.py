
from docx import Document
from docx.shared import Pt
import datetime

def main():
    doc = Document()
    
    # Title Page for Pathogens & Immunity
    # Title
    p = doc.add_paragraph("Molecular Determinants of BCG Vaccine Response Heterogeneity: A Systematic Review and Meta-Analysis of Multi-Omics Data")
    p.style.font.name = 'Times New Roman'
    p.style.font.size = Pt(14)
    p.style.font.bold = True
    
    # Author
    p = doc.add_paragraph("\nSiddalingaiah H S, MD")
    p.style.font.name = 'Times New Roman'
    p.style.font.size = Pt(12)
    p.style.font.bold = True
    
    # Affiliation
    p = doc.add_paragraph("Professor, Department of Community Medicine\nShridevi Institute of Medical Sciences and Research Hospital\nTumkur, Karnataka, India")
    p.style.font.name = 'Times New Roman'
    p.style.font.size = Pt(12)
    
    # Corresponding Author
    p = doc.add_paragraph("\nCorresponding Author:")
    p.style.font.name = 'Times New Roman'
    p.style.font.size = Pt(12)
    p.style.font.bold = True
    
    p = doc.add_paragraph("Siddalingaiah H S, MD\nShridevi Institute of Medical Sciences and Research Hospital\nSira Road, Tumkur - 572106, Karnataka, India\nEmail: hssling@yahoo.com")
    p.style.font.name = 'Times New Roman'
    p.style.font.size = Pt(12)
    
    # Running Head
    p = doc.add_paragraph("\nRunning Head: Meta-Analysis of BCG Response Heterogeneity")
    p.style.font.name = 'Times New Roman'
    p.style.font.size = Pt(12)
    
    # Keywords
    p = doc.add_paragraph("Keywords: BCG vaccine, Trained immunity, Epigenetics, Meta-analysis, Personalized medicine")
    
    # Word Count
    p = doc.add_paragraph("\nWord Count: ~3500 words")
    p = doc.add_paragraph("Number of Figures: 2")
    p = doc.add_paragraph("Number of Tables: 2")
    
    # Declarations
    p = doc.add_paragraph("\n\nDeclarations")
    p.style.font.bold = True
    
    p = doc.add_paragraph("Funding: None")
    p = doc.add_paragraph("Conflicts of Interest: None declared")
    p = doc.add_paragraph("Data Availability: All data analyzed are publicly available. Source code is archived at: https://github.com/hssling/BCG_Vaccine_Response")
    p = doc.add_paragraph("Authors contributions: SHS conceived the study, analyzed data, and wrote the manuscript.")
    
    doc.save("d:/research-automation/TB multiomics/TB Chromatin Priming Multiomics/BCG_Vaccine_Response/4_manuscript/TitlePage_PathogensImmunity.docx")
    print("Title Page generated.")

if __name__ == "__main__":
    main()
