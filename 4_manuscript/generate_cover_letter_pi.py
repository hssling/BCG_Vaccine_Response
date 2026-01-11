
from docx import Document
from docx.shared import Pt
import datetime

def main():
    doc = Document()
    today = datetime.date.today().strftime("%B %d, %Y")
    
    # Header
    p = doc.add_paragraph(f"{today}\n")
    p = doc.add_paragraph("Dr. Michael Lederman")
    p = doc.add_paragraph("Editor-in-Chief")
    p = doc.add_paragraph("Pathogens & Immunity")
    
    # Salutation
    p = doc.add_paragraph("\nDear Dr. Lederman,")
    
    # Opening
    text = (
        "I am pleased to submit my manuscript titled \"Molecular Determinants of BCG Vaccine Response Heterogeneity: "
        "A Systematic Review and Meta-Analysis of Multi-Omics Data\" for consideration as an Original Research Article "
        "in Pathogens & Immunity."
    )
    doc.add_paragraph(text)
    
    # The "Why P&I" paragraph
    text = (
        "Given Pathogens & Immunity's focus on microbial pathogenesis and host defense, particularly its leadership in publishing "
        "cutting-edge research on trained immunity (e.g., recent work by Netea et al.), I believe this work is ideally suited for "
        "your readership. While P&I typically does not accept unsolicited reviews, this manuscript presents an original quantitative "
        "meta-analysis that synthesizes multi-omics data for the first time, generating new pooled estimates (Ratio: 1.96, I²: 0%) "
        "that validate IL-1β as a robust predictor of vaccine response."
    )
    doc.add_paragraph(text)
    
    # The "What we found" paragraph
    text = (
        "BCG is the world's most widely used vaccine, yet efficacy is highly variable. "
        "My analysis of 361 individuals across multiple independent cohorts reveals a consistent trimodal response distribution. "
        "Crucially, I demonstrate that baseline IL-1β production capacity is not just a correlate but a validated predictor of "
        "trained immunity induction, with zero heterogeneity across studies. These findings provide the first quantitative "
        "rationale for stratified vaccination strategies."
    )
    doc.add_paragraph(text)
    
    # Closing
    text = (
        "This manuscript has not been published and is not under consideration for publication elsewhere. "
        "I have no conflicts of interest to disclose.\n\n"
        "Thank you for considering this submission. I look forward to your response."
    )
    doc.add_paragraph(text)
    
    # Sign-off
    p = doc.add_paragraph("\nSincerely,")
    p = doc.add_paragraph("Siddalingaiah H S, MD")
    p = doc.add_paragraph("Professor, Department of Community Medicine")
    p = doc.add_paragraph("Shridevi Institute of Medical Sciences and Research Hospital")
    p = doc.add_paragraph("Tumkur, Karnataka, India")
    p = doc.add_paragraph("Email: hssling@yahoo.com")
    
    doc.save("d:/research-automation/TB multiomics/TB Chromatin Priming Multiomics/BCG_Vaccine_Response/4_manuscript/Cover_Letter_PathogensImmunity.docx")
    print("Cover Letter generated.")

if __name__ == "__main__":
    main()
