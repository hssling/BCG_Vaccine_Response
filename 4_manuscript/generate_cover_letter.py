
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def create_cover_letter():
    doc = Document()
    
    # Header material
    today = datetime.now().strftime("%B %d, %Y")
    doc.add_paragraph(f"{today}")
    doc.add_paragraph("Editorial Office\nPathogens & Immunity\n")
    
    doc.add_paragraph("RE: Submission of Original Meta-Analysis Research entitled \"Molecular Determinants of BCG Vaccine Response Heterogeneity: A Systematic Review and Meta-Analysis of Multi-Omics Data\"")
    
    doc.add_paragraph("Dear Editor-in-Chief,")
    
    doc.add_paragraph(
        "We are pleased to submit our Original Meta-Analysis for consideration in Pathogens & Immunity. "
        "While we understand the journal typically solicits narrative reviews, we present here a quantitative "
        "Original Research study utilizing secondary data analysis methods (PRISMA guidelines). "
        "The Bacillus Calmette-Guérin (BCG) vaccine is the most widely used vaccine globally, yet its efficacy "
        "varies dramatically (0-80%)."
    )
    
    doc.add_paragraph(
        "In this study, we systematically synthesized multi-omics data from six major studies (published 2012-2024), "
        "comprising a total of 662 individuals. By integrating findings from single-cell RNA sequencing, DNA methylation profiling, "
        "and histone modification mapping, we identified robust molecular predictors of vaccine response."
    )
    
    doc.add_paragraph(
        "Key findings include:"
    )
    
    # Bullet points
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.add_run("Trimodal Response Distribution: ").bold = True
    p.add_run("We confirmed a consistent phenotype where ~30% of individuals are high responders, driven by baseline epigenetic states.")
    
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.add_run("Epigenetic Priming: ").bold = True
    p.add_run("Consolidated evidence shows that baseline H3K4me3 enrichment at inflammatory promoters (OR 2.8) is the strongest predictor of trained immunity induction.")
    
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.add_run("Long-term Persistence: ").bold = True
    p.add_run("Our review confirms that BCG-induced DNA methylation signatures persist in monocytes for over 12 months (Bannister et al., 2022).")
    
    doc.add_paragraph(
        "This manuscript provides the first quantitative synthesis of these disparate multi-omics datasets, creating a "
        "unified framework for understanding vaccine response heterogeneity. We believe this work will be of broad interest "
        "to immunologists, vaccinologists, and clinicians seeking to optimize BCG usage."
    )
    
    doc.add_paragraph("We confirm that this manuscript is original, has not been published elsewhere, and is not under consideration by another journal. The author declares no conflicts of interest.")
    
    doc.add_paragraph("Thank you for your consideration.")
    
    doc.add_paragraph("Sincerely,\n\nSiddalingaiah H S, MD\nProfessor, Department of Community Medicine\nShridevi Institute of Medical Sciences\nTumkur, Karnataka, India\nEmail: hssling@yahoo.com")
    
    doc.save("Cover_Letter_BCG_Systematic_Review.docx")
    print("Cover letter saved.")

if __name__ == "__main__":
    create_cover_letter()
