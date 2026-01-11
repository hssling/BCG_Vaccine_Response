
from docx import Document
from docx.shared import Inches, Pt
import os
from pathlib import Path

BASE_DIR = Path("d:/research-automation/TB multiomics/TB Chromatin Priming Multiomics/BCG_Vaccine_Response")
FIG_DIR = BASE_DIR / "3_results" / "figures"
OUTPUT_DIR = BASE_DIR / "4_manuscript"

def main():
    doc = Document()
    
    # Title
    p = doc.add_paragraph("Supplementary Materials")
    p.style.font.name = 'Times New Roman'
    p.style.font.size = Pt(16)
    p.style.font.bold = True
    p.alignment = 1 # Center
    
    # Subtitle
    p = doc.add_paragraph("Molecular Determinants of BCG Vaccine Response Heterogeneity: A Systematic Review and Meta-Analysis")
    p.style.font.size = Pt(12)
    p.alignment = 1
    
    doc.add_page_break()
    
    # Figure S1
    doc.add_heading("Supplementary Figures", 1)
    
    if os.path.exists(FIG_DIR / "FigS1_PRISMA_Flow_Diagram.png"):
        doc.add_picture(str(FIG_DIR / "FigS1_PRISMA_Flow_Diagram.png"), width=Inches(6))
        
        # Caption
        p = doc.add_paragraph("Figure S1. PRISMA 2020 flow diagram for the systematic review.")
        p.style.font.bold = True
        
        caption_text = (
            "The flow diagram depicts the flow of information through the different phases of the systematic review. "
            "It maps out the number of records identified, included and excluded, and the reasons for exclusions."
        )
        doc.add_paragraph(caption_text)
    else:
        doc.add_paragraph("[Figure S1 Placeholder - Image not found]")

    doc.add_page_break()
    
    # Supplementary Methods - Search Strategy
    doc.add_heading("Supplementary Note 1: Detailed Search Strategy", 1)
    
    doc.add_paragraph("Database: PubMed/MEDLINE")
    doc.add_paragraph("Date of Search: December 31, 2024")
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Step'
    hdr[1].text = 'Search Query'
    
    queries = [
        ("1", "BCG Vaccine OR Bacillus Calmette-Guerin OR trained immunity"),
        ("2", "heterogeneity OR variability OR responder OR non-responder OR variation"),
        ("3", "transcriptomics OR RNA-seq OR scRNA-seq OR microarray OR gene expression"),
        ("4", "epigenetics OR DNA methylation OR histone modification OR ATAC-seq OR ChIP-seq"),
        ("5", "metabolomics OR lipidomics OR metabolic profiling"),
        ("6", "3 OR 4 OR 5"),
        ("7", "1 AND 2 AND 6"),
        ("8", "Limit to Humans, English Language, 2012-2024")
    ]
    
    for step, q in queries:
        row = table.add_row().cells
        row[0].text = step
        row[1].text = q
        
    output_path = OUTPUT_DIR / "Supplementary_Materials.docx"
    doc.save(output_path)
    print(f"Supplementary Materials saved to {output_path}")

if __name__ == "__main__":
    main()
